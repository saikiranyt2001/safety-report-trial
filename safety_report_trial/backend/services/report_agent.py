# Report Agent for Structured Safety Reports

from docx import Document
from backend.rag.rag_engine import retrieve_regulation

def generate_structured_report(site_data, hazards, risks, controls, compliance):
    report = f"""
INDUSTRIAL SAFETY REPORT

Company: {site_data.get('company', '')}
Location: {site_data.get('location', '')}
Inspection Date: {site_data.get('inspection_date', '')}

Hazards Identified:
{chr(10).join([f"• {h}" for h in hazards])}

Risk Assessment:
{chr(10).join([f"{h} → {risks.get(h, '')}" for h in hazards])}

Control Measures:
{chr(10).join([f"{h}: {controls.get(h, '')}" for h in hazards])}

Compliance:
{chr(10).join([f"- {c}" for c in compliance])}

Regulation Evidence:
{chr(10).join([f"{h}: {chr(10).join(retrieve_regulation(h))}" for h in hazards])}

Recommendations:
{site_data.get('recommendations', '')}

Review Date: {site_data.get('review_date', '')}
"""
    return report


def export_report(report_text, filename="report.docx"):
    doc = Document()
    doc.add_heading("Industrial Safety Report", level=1)
    for line in report_text.split("\n"):
        doc.add_paragraph(line)
    doc.save(filename)
