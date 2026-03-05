# export_docx.py
# Export to DOCX logic

from docx import Document

def export_docx(content):

    doc = Document()

    doc.add_heading("Safety Report", level=1)

    for line in content.split("\n"):
        doc.add_paragraph(line)

    doc.save("report.docx")